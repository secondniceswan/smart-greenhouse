from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont, ImageFilter, ImageEnhance
import os, math

OUT_DIR = '/home/ninalover/Documents/prd'
IMG_DIR = os.path.join(OUT_DIR, '_img')
REAL_DIR = os.path.join(IMG_DIR, 'real')
os.makedirs(IMG_DIR, exist_ok=True)

# ============================================================
# FONT HELPER
# ============================================================
_font_cache = {}
def get_font(size=16, bold=False):
    key = (size, bold)
    if key in _font_cache:
        return _font_cache[key]
    paths = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf' if bold else '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/dejavu-sans-fonts/DejaVuSans-Bold.ttf' if bold else '/usr/share/fonts/dejavu-sans-fonts/DejaVuSans.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf' if bold else '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
        '/usr/share/fonts/google-noto/NotoSans-Bold.ttf' if bold else '/usr/share/fonts/google-noto/NotoSans-Regular.ttf',
    ]
    for p in paths:
        if os.path.exists(p):
            f = ImageFont.truetype(p, size)
            _font_cache[key] = f
            return f
    try:
        f = ImageFont.truetype("DejaVuSans.ttf", size)
    except:
        f = ImageFont.load_default()
    _font_cache[key] = f
    return f

def get_mono(size=14):
    mono_paths = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf',
        '/usr/share/fonts/dejavu-sans-mono-fonts/DejaVuSansMono.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationMono-Regular.ttf',
    ]
    for p in mono_paths:
        if os.path.exists(p):
            return ImageFont.truetype(p, size)
    return get_font(size)

# ============================================================
# COLOR PALETTE
# ============================================================
BG          = (22, 22, 38)
BG_CARD     = (32, 32, 52)
BG_DARK     = (16, 16, 28)
WHITE       = (240, 240, 245)
LIGHT_GRAY  = (180, 180, 200)
GRAY        = (120, 120, 145)
DARK_GRAY   = (70, 70, 90)
ACCENT      = (78, 205, 196)
ACCENT_DIM  = (50, 140, 133)
RED         = (255, 90, 90)
RED_DIM     = (160, 50, 50)
ORANGE      = (255, 170, 70)
GREEN       = (100, 220, 120)
GREEN_DIM   = (50, 120, 60)
BLUE        = (100, 150, 255)
YELLOW      = (240, 230, 80)
PURPLE      = (190, 120, 255)
PINK        = (255, 130, 170)
WIRE_RED    = (220, 50, 50)
WIRE_BLUE   = (50, 100, 220)
WIRE_ORANGE = (230, 140, 30)
WIRE_BROWN  = (150, 90, 40)
WIRE_GREEN  = (50, 180, 80)

# ============================================================
# LOAD REAL PHOTOS
# ============================================================
def load_real(filename, max_w=None, max_h=None):
    """Load a real photo, resize proportionally, return as RGBA Image."""
    path = os.path.join(REAL_DIR, filename)
    if not os.path.exists(path):
        # Create placeholder
        img = Image.new('RGBA', (max_w or 200, max_h or 150), (60, 60, 80, 255))
        d = ImageDraw.Draw(img)
        d.text((10, 10), f"[{filename}]", fill=WHITE, font=get_font(12))
        return img
    img = Image.open(path).convert('RGBA')
    if max_w and max_h:
        img.thumbnail((max_w, max_h), Image.LANCZOS)
    elif max_w:
        ratio = max_w / img.width
        img = img.resize((max_w, int(img.height * ratio)), Image.LANCZOS)
    elif max_h:
        ratio = max_h / img.height
        img = img.resize((int(img.width * ratio), max_h), Image.LANCZOS)
    return img

def paste_real(canvas, real_img, x, y):
    """Paste a real photo onto canvas at (x, y)."""
    canvas.paste(real_img, (x, y), real_img if real_img.mode == 'RGBA' else None)

def add_shadow_border(img, border=3, shadow=6, border_color=(200,200,200,255)):
    """Add a border and drop shadow to a component photo."""
    w, h = img.size
    new_w = w + border*2 + shadow
    new_h = h + border*2 + shadow
    result = Image.new('RGBA', (new_w, new_h), (0,0,0,0))
    # Shadow
    shadow_img = Image.new('RGBA', (w+border*2, h+border*2), (0,0,0,80))
    result.paste(shadow_img, (shadow, shadow))
    # Border
    border_img = Image.new('RGBA', (w+border*2, h+border*2), border_color)
    result.paste(border_img, (0, 0))
    # Image
    result.paste(img, (border, border), img if img.mode == 'RGBA' else None)
    return result

# ============================================================
# DRAWING PRIMITIVES
# ============================================================
def make_img(name, w, h, draw_func):
    img = Image.new('RGB', (w, h), BG)
    d = ImageDraw.Draw(img)
    draw_func(img, d, w, h)  # Now pass img too for compositing
    path = os.path.join(IMG_DIR, f'{name}.png')
    img.save(path, quality=95)
    return path

def draw_title_bar(d, w, text, subtitle=None):
    d.rectangle([0, 0, w, 50], fill=(28, 28, 48))
    d.line([(0, 50), (w, 50)], fill=ACCENT, width=2)
    d.text((20, 10), text, fill=ACCENT, font=get_font(18, bold=True))
    if subtitle:
        d.text((20, 32), subtitle, fill=GRAY, font=get_font(11))

def draw_wire(d, points, color, width=3, dashed=False):
    for i in range(len(points)-1):
        if dashed:
            x1, y1 = points[i]
            x2, y2 = points[i+1]
            length = math.sqrt((x2-x1)**2 + (y2-y1)**2)
            dashes = max(1, int(length / 12))
            for j in range(dashes):
                t1 = j / dashes
                t2 = min(1, (j + 0.6) / dashes)
                sx = x1 + (x2-x1)*t1; sy = y1 + (y2-y1)*t1
                ex = x1 + (x2-x1)*t2; ey = y1 + (y2-y1)*t2
                d.line([(sx,sy),(ex,ey)], fill=color, width=width)
        else:
            d.line([points[i], points[i+1]], fill=color, width=width)

def draw_arrow(d, x1, y1, x2, y2, color, width=3, head=12):
    d.line([(x1,y1),(x2,y2)], fill=color, width=width)
    angle = math.atan2(y2-y1, x2-x1)
    a1 = angle + math.pi * 0.8
    a2 = angle - math.pi * 0.8
    d.polygon([(x2,y2),(x2+head*math.cos(a1),y2+head*math.sin(a1)),(x2+head*math.cos(a2),y2+head*math.sin(a2))], fill=color)

def draw_warning_box(d, x, y, w, text, subtext=None):
    h = 50 if not subtext else 70
    d.rounded_rectangle([x, y, x+w, y+h], radius=8, fill=(80, 20, 20), outline=RED, width=2)
    cx = x + 20; cy = y + 16
    d.polygon([(cx, cy-8), (cx-9, cy+7), (cx+9, cy+7)], fill=YELLOW, outline=RED)
    d.text((cx-2, cy-5), "!", fill=(40,0,0), font=get_font(10, True))
    d.text((x+38, y+10), text, fill=RED, font=get_font(12, True))
    if subtext:
        d.text((x+38, y+32), subtext, fill=ORANGE, font=get_font(11))

def draw_info_box(d, x, y, w, text, subtext=None):
    h = 50 if not subtext else 70
    d.rounded_rectangle([x, y, x+w, y+h], radius=8, fill=(20, 40, 60), outline=BLUE, width=2)
    d.text((x+15, y+10), text, fill=BLUE, font=get_font(12, True))
    if subtext:
        d.text((x+15, y+32), subtext, fill=LIGHT_GRAY, font=get_font(11))

def draw_label_line(d, x1, y1, x2, y2, label, color, label_side="end"):
    """Draw a line with a label at the end/start."""
    d.line([(x1,y1),(x2,y2)], fill=color, width=2)
    # Dot at connection
    d.ellipse([x1-4, y1-4, x1+4, y1+4], fill=color)
    d.ellipse([x2-4, y2-4, x2+4, y2+4], fill=color)
    f = get_font(10, True)
    if label_side == "end":
        d.text((x2+8, y2-7), label, fill=color, font=f)
    else:
        bb = d.textbbox((0,0), label, font=f)
        d.text((x1-8-(bb[2]-bb[0]), y1-7), label, fill=color, font=f)

def draw_legend(d, x, y, items):
    d.rounded_rectangle([x-10, y-10, x+230, y + len(items)*24 + 10], radius=6, fill=BG_CARD, outline=DARK_GRAY)
    d.text((x, y-6), "LEGENDA KABEL", fill=GRAY, font=get_font(9, True))
    y += 14
    for color, label in items:
        d.line([(x, y+7), (x+24, y+7)], fill=color, width=4)
        d.ellipse([x+20, y+3, x+28, y+11], fill=color)
        d.text((x+32, y), label, fill=LIGHT_GRAY, font=get_font(10))
        y += 24

def draw_pin_label(d, x, y, pin_name, target, color, direction="right"):
    """Draw a labeled pin connection: pin_name --> target"""
    f = get_font(11, True)
    fs = get_font(10)
    if direction == "right":
        d.ellipse([x-5, y-5, x+5, y+5], fill=color, outline=WHITE)
        d.text((x+10, y-8), pin_name, fill=color, font=f)
        d.text((x+10, y+6), target, fill=GRAY, font=fs)
    else:
        d.ellipse([x-5, y-5, x+5, y+5], fill=color, outline=WHITE)
        bb = d.textbbox((0,0), pin_name, font=f)
        d.text((x-12-(bb[2]-bb[0]), y-8), pin_name, fill=color, font=f)
        bb2 = d.textbbox((0,0), target, font=fs)
        d.text((x-12-(bb2[2]-bb2[0]), y+6), target, fill=GRAY, font=fs)

# ============================================================
# IMAGE GENERATORS WITH REAL PHOTOS
# ============================================================

def img_usb_connect(img, d, w, h):
    draw_title_bar(d, w, "Colok ESP32-S3 ke Laptop via USB-C", "Koneksi pertama: ESP32 ke laptop dengan kabel USB-C DATA")

    # Laptop (drawn)
    lx, ly = 30, 80
    d.rounded_rectangle([lx, ly, lx+280, ly+190], radius=8, fill=(45,45,65), outline=(80,80,100), width=2)
    d.rounded_rectangle([lx+10, ly+10, lx+270, ly+170], radius=4, fill=(30,30,48))
    d.text((lx+20, ly+25), "Device Manager", fill=ACCENT, font=get_font(13, True))
    d.text((lx+20, ly+50), "Ports (COM & LPT)", fill=WHITE, font=get_font(11))
    d.text((lx+40, ly+72), "COM3 — USB Serial", fill=GREEN, font=get_font(12, True))
    d.text((lx+40, ly+92), "(ESP32-S3 detected)", fill=GREEN, font=get_font(10))
    d.text((lx+20, ly+120), "Kalau ini MUNCUL =", fill=GRAY, font=get_font(10))
    d.text((lx+20, ly+136), "kabel USB BENAR", fill=GREEN, font=get_font(12, True))
    d.rounded_rectangle([lx-10, ly+190, lx+290, ly+215], radius=4, fill=(55,55,75), outline=(80,80,100))

    # USB Cable illustration
    d.text((330, 120), "KABEL", fill=ORANGE, font=get_font(14, True))
    d.text((330, 140), "USB-C", fill=ORANGE, font=get_font(14, True))
    d.text((330, 162), "DATA", fill=WHITE, font=get_font(12, True))
    draw_wire(d, [(310, 180), (350, 180), (430, 160), (470, 160)], ORANGE, 6)
    d.text((330, 190), "USB-A", fill=GRAY, font=get_font(9))
    d.text((460, 145), "USB-C", fill=GRAY, font=get_font(9))

    # ESP32 REAL PHOTO
    esp_photo = load_real('esp32s3.png', max_h=230)
    esp_photo = add_shadow_border(esp_photo)
    paste_real(img, esp_photo, 490, 70)
    d = ImageDraw.Draw(img)  # Refresh draw context
    # Label
    d.rounded_rectangle([490, 310, 730, 345], radius=6, fill=BG_CARD, outline=GREEN)
    d.text((500, 315), "ESP32-S3-N16R8 Soldered", fill=GREEN, font=get_font(12, True))

    # Warning and info boxes
    draw_warning_box(d, 30, h-130, 370, "Kabel HARUS USB DATA!", "Kabel charge-only TIDAK punya jalur data D+/D-")
    draw_info_box(d, 420, h-130, 370, "Cara cek: Device Manager", "Harus muncul COM port baru (COM3/COM4)")

    # Data vs charge comparison
    d.rounded_rectangle([30, h-48, 790, h-8], radius=6, fill=BG_CARD)
    d.text((45, h-44), "Data cable:", fill=GREEN, font=get_font(10, True))
    d.text((145, h-44), "4 kabel internal (VCC, D+, D-, GND) — bisa transfer data", fill=LIGHT_GRAY, font=get_font(10))
    d.text((45, h-26), "Charge-only:", fill=RED, font=get_font(10, True))
    d.text((145, h-26), "2 kabel internal (VCC, GND saja) — TIDAK bisa transfer data", fill=LIGHT_GRAY, font=get_font(10))

def img_arduino_board(img, d, w, h):
    draw_title_bar(d, w, "Setting Board di Arduino IDE", "Tools > Board Settings — setiap setting yang merah = KRITIS")

    ide_x, ide_y = 30, 65
    ide_w = w - 60
    d.rounded_rectangle([ide_x, ide_y, ide_x+ide_w, ide_y+30], radius=8, fill=(50,50,70))
    d.text((ide_x+15, ide_y+8), "Arduino IDE 2.x — Tools > Board Configuration", fill=LIGHT_GRAY, font=get_font(11))
    for i, c in enumerate([RED, YELLOW, GREEN]):
        d.ellipse([ide_x+ide_w-50+i*16, ide_y+8, ide_x+ide_w-38+i*16, ide_y+20], fill=c)

    settings = [
        ("Board", "ESP32S3 Dev Module", GREEN, True, "Board manager ESP32 by Espressif"),
        ("PSRAM", "OPI PSRAM", RED, True, "KRITIS! Harus OPI, BUKAN QSPI!"),
        ("Flash Size", "16MB (128Mb)", WHITE, False, "Sesuai spek board N16R8"),
        ("USB CDC On Boot", "Enabled", RED, True, "KRITIS! Tanpa ini Serial Monitor MATI"),
        ("Upload Speed", "921600", LIGHT_GRAY, False, "Speed tercepat"),
        ("Partition Scheme", "Default 4MB with spiffs", LIGHT_GRAY, False, "Default cukup"),
        ("Port", "COM3 (atau yg muncul)", ORANGE, False, "Pilih COM yg baru muncul"),
    ]

    y = ide_y + 38
    for label, value, color, critical, note in settings:
        row_bg = (55, 30, 30) if critical else (40, 40, 58)
        d.rounded_rectangle([ide_x+5, y, ide_x+ide_w-5, y+44], radius=4, fill=row_bg)
        if critical:
            d.rounded_rectangle([ide_x+9, y+4, ide_x+13, y+40], radius=2, fill=RED)
        d.text((ide_x+20, y+5), label, fill=GRAY, font=get_font(12))
        val_x = ide_x + 210
        d.rounded_rectangle([val_x, y+4, val_x+300, y+26], radius=4, fill=BG_DARK, outline=(80,80,100))
        d.text((val_x+8, y+6), value, fill=color, font=get_font(12, True))
        d.text((ide_x+20, y+28), note, fill=DARK_GRAY, font=get_font(9))
        y += 50

    # ESP32 photo on the right as reference
    esp_photo = load_real('ESP32-S3-DevKitC-1_v2-annotated-photo.png', max_w=270)
    paste_real(img, esp_photo, w-290, ide_y+50)
    d = ImageDraw.Draw(img)
    d.text((w-290, ide_y+40), "Board yang dipakai:", fill=GRAY, font=get_font(9))

    draw_warning_box(d, 30, h-65, w-60, "PSRAM = OPI & USB CDC = Enabled WAJIB BENAR!", "Salah setting: board tidak terdeteksi atau PSRAM out of memory")

def img_libraries(img, d, w, h):
    draw_title_bar(d, w, "Install 5 Library via Library Manager", "Tools > Manage Libraries > cari & install satu per satu")

    libs = [
        ("BH1750", "Christopher Laws", "1.3.0", "Sensor cahaya I2C (lux)", ACCENT),
        ("DHT sensor library", "Adafruit", "1.4.6", "Sensor suhu & humidity DHT11", GREEN),
        ("Adafruit Unified Sensor", "Adafruit", "1.1.14", "WAJIB! Dependency dari DHT library", ORANGE),
        ("ESP32Servo", "Kevin Harrington", "3.0.5", "Kontrol servo motor PWM di ESP32", PURPLE),
        ("ArduinoJson", "Benoit Blanchon", "7.x", "Parsing JSON dari Supabase & Open-Meteo", BLUE),
    ]

    y = 68
    for i, (name, author, ver, desc, color) in enumerate(libs):
        d.rounded_rectangle([35, y, w-35, y+68], radius=8, fill=BG_CARD, outline=DARK_GRAY)
        d.rounded_rectangle([45, y+8, 75, y+38], radius=14, fill=color)
        d.text((55, y+12), str(i+1), fill=BG, font=get_font(14, True))
        d.text((85, y+8), name, fill=color, font=get_font(14, True))
        d.text((85, y+28), f"by {author}  |  v{ver}", fill=GRAY, font=get_font(10))
        d.text((85, y+44), desc, fill=LIGHT_GRAY, font=get_font(10))
        btn_x = w - 145
        d.rounded_rectangle([btn_x, y+15, btn_x+95, y+45], radius=6, fill=GREEN_DIM, outline=GREEN)
        d.text((btn_x+15, y+20), "INSTALL", fill=WHITE, font=get_font(11, True))
        y += 78

    draw_info_box(d, 35, h-55, w-70, "Pastikan SEMUA 5 library INSTALLED sebelum lanjut ke compile!", None)

def img_breadboard_esp(img, d, w, h):
    draw_title_bar(d, w, "Pasang ESP32-S3 di Breadboard 830 Point", "Tancapkan di tengah — kedua sisi pin harus bisa diakses")

    # ESP32 pinout photo (real)
    pinout = load_real('ESP32-S3_DevKitC-1_pinlayout_v1.1.jpg', max_w=700)
    pinout = add_shadow_border(pinout, border=4, shadow=8)
    paste_real(img, pinout, (w - pinout.width) // 2, 65)
    d = ImageDraw.Draw(img)

    pin_y = 65 + pinout.height + 15

    # Key pin labels we use
    d.rounded_rectangle([30, pin_y, w-30, pin_y+120], radius=8, fill=BG_CARD, outline=DARK_GRAY)
    d.text((50, pin_y+8), "PIN YANG KITA PAKAI:", fill=ACCENT, font=get_font(13, True))

    pins_used = [
        ("GPIO8", "SDA (I2C data) -- ke BH1750", ACCENT),
        ("GPIO9", "SCL (I2C clock) -- ke BH1750", ACCENT),
        ("GPIO4", "DATA DHT11 (suhu/humidity)", GREEN),
        ("GPIO5", "DO Rain Sensor (hujan)", YELLOW),
        ("GPIO13", "Signal Servo MG996R", PURPLE),
        ("GPIO2", "LED Warning (anti-oven)", PINK),
        ("VIN", "Power input 5V dari adaptor", WIRE_RED),
        ("3V3", "Output 3.3V untuk sensor", ORANGE),
        ("GND", "Ground (bersama semua)", WIRE_BLUE),
    ]
    col1 = pins_used[:5]
    col2 = pins_used[5:]
    for i, (pin, desc, color) in enumerate(col1):
        dy = pin_y + 30 + i * 17
        d.text((50, dy), pin, fill=color, font=get_font(10, True))
        d.text((120, dy), desc, fill=LIGHT_GRAY, font=get_font(9))
    for i, (pin, desc, color) in enumerate(col2):
        dy = pin_y + 30 + i * 17
        d.text((420, dy), pin, fill=color, font=get_font(10, True))
        d.text((470, dy), desc, fill=LIGHT_GRAY, font=get_font(9))

    draw_info_box(d, 30, h-55, w-60, "Tekan ESP32 kuat — semua pin masuk lubang. Kedua sisi harus bisa dipasangi jumper.", None)

def img_power_wiring(img, d, w, h):
    draw_title_bar(d, w, "Wiring Power Supply (5V 3A Adaptor)", "Adaptor 220V AC > Jack DC > Breadboard Rail > ESP32 VIN")

    # Adaptor (drawn - hard to find free photo)
    d.rounded_rectangle([33, 83, 193, 203], radius=8, fill=(10,10,20))
    d.rounded_rectangle([30, 80, 190, 200], radius=8, fill=(40,40,40), outline=(80,80,80), width=2)
    d.text((45, 90), "ADAPTOR", fill=WHITE, font=get_font(14, True))
    d.text((45, 112), "Input:", fill=GRAY, font=get_font(10))
    d.text((90, 112), "AC 220V", fill=LIGHT_GRAY, font=get_font(10))
    d.text((45, 130), "Output:", fill=GRAY, font=get_font(10))
    d.text((100, 130), "DC 5V 3A", fill=RED, font=get_font(12, True))
    d.text((45, 152), "Jack: 5.5x2.1mm", fill=GRAY, font=get_font(10))
    d.text((45, 170), "Ke stopkontak", fill=DARK_GRAY, font=get_font(9))
    # DC plug
    d.ellipse([178, 120, 200, 150], fill=(30,30,30), outline=(100,100,100), width=2)

    # Cable to jack DC
    draw_wire(d, [(200, 135), (250, 135)], (60,60,60), 7)
    d.text((210, 115), "kabel DC", fill=GRAY, font=get_font(9))

    # Jack DC screw terminal (drawn)
    jx = 255
    d.rounded_rectangle([jx+3, 103, jx+163, 173], radius=6, fill=(10,10,20))
    d.rounded_rectangle([jx, 100, jx+160, 170], radius=6, fill=(50,50,35), outline=(100,100,70), width=2)
    d.text((jx+10, 105), "JACK DC FEMALE", fill=WHITE, font=get_font(11, True))
    d.text((jx+10, 122), "Screw Terminal", fill=LIGHT_GRAY, font=get_font(10))
    d.rounded_rectangle([jx+10, 140, jx+70, 165], radius=4, fill=(30,30,30), outline=WIRE_RED, width=2)
    d.text((jx+18, 144), "+ (5V)", fill=RED, font=get_font(10, True))
    d.rounded_rectangle([jx+80, 140, jx+150, 165], radius=4, fill=(30,30,30), outline=WIRE_BLUE, width=2)
    d.text((jx+88, 144), "- (GND)", fill=BLUE, font=get_font(10, True))

    # ESP32 photo
    esp_photo = load_real('esp32s3.png', max_h=180)
    esp_photo = add_shadow_border(esp_photo, border=3, shadow=5)
    paste_real(img, esp_photo, 560, 70)
    d = ImageDraw.Draw(img)
    d.text((570, 255), "ESP32-S3", fill=GREEN, font=get_font(10, True))

    # Wires from jack to ESP32
    # Red: + to VIN
    draw_wire(d, [(jx+40, 140), (jx+40, 80), (580, 80), (580, 100)], WIRE_RED, 4)
    d.rounded_rectangle([440, 68, 560, 92], radius=4, fill=BG_CARD, outline=WIRE_RED)
    d.text((448, 72), "MERAH: +5V ke VIN", fill=RED, font=get_font(10, True))

    # Blue: - to GND
    draw_wire(d, [(jx+115, 140), (jx+115, 65), (650, 65), (650, 100)], WIRE_BLUE, 4)
    d.rounded_rectangle([440, 48, 560, 68], radius=4, fill=BG_CARD, outline=WIRE_BLUE)
    d.text((448, 50), "HITAM: GND", fill=BLUE, font=get_font(10, True))

    # Connection table
    y_info = 280
    d.rounded_rectangle([30, y_info, w-30, y_info+155], radius=8, fill=BG_CARD, outline=DARK_GRAY)
    d.text((50, y_info+8), "URUTAN KONEKSI POWER:", fill=WHITE, font=get_font(14, True))

    steps = [
        ("1.", "Jack DC (+) ---------> Power Rail +5V breadboard", WIRE_RED, "kabel merah"),
        ("2.", "Jack DC (-) ---------> Ground Rail GND breadboard", WIRE_BLUE, "kabel hitam"),
        ("3.", "Power Rail (+5V) ----> ESP32 pin VIN", WIRE_RED, "jumper merah"),
        ("4.", "Ground Rail (GND) --> ESP32 pin GND", WIRE_BLUE, "jumper hitam"),
        ("5.", "ESP32 pin 3.3V -----> Baris kosong breadboard", ORANGE, "jadi 3.3V bus sensor!"),
    ]
    for i, (num, text, color, note) in enumerate(steps):
        sy = y_info + 35 + i * 22
        d.text((50, sy), num, fill=color, font=get_font(11, True))
        d.text((75, sy), text, fill=WHITE, font=get_font(11))
        d.text((560, sy), note, fill=GRAY, font=get_font(9))

    draw_warning_box(d, 30, h-68, w-60, "VIN ESP32 = 5V. Sensor = 3.3V. JANGAN TUKAR!", "5V langsung ke sensor = sensor RUSAK permanen!")

def img_sensor_wiring(img, d, w, h):
    draw_title_bar(d, w, "Wiring 3 Sensor ke ESP32-S3", "BH1750 (cahaya) + DHT11 (suhu) + Rain Sensor (hujan)")

    # ESP32 photo (center reference)
    esp_photo = load_real('esp32s3.png', max_h=180)
    esp_photo = add_shadow_border(esp_photo, border=2, shadow=4)
    esp_cx = 30
    esp_cy = 80
    paste_real(img, esp_photo, esp_cx, esp_cy)
    d = ImageDraw.Draw(img)
    d.text((esp_cx, esp_cy + esp_photo.height + 5), "ESP32-S3", fill=GREEN, font=get_font(10, True))

    # Pin connection points on ESP32 (approximate positions on right side)
    esp_right = esp_cx + esp_photo.width
    pin_3v3_y = esp_cy + 30
    pin_gnd_y = esp_cy + 55
    pin_8_y   = esp_cy + 85
    pin_9_y   = esp_cy + 105
    pin_4_y   = esp_cy + 130
    pin_5_y   = esp_cy + 155

    # ====== BH1750 REAL PHOTO ======
    bh_photo = load_real('bh1750_real.jpg', max_h=140)
    bh_photo = add_shadow_border(bh_photo, border=2, shadow=4)
    bh_x, bh_y = 380, 65
    paste_real(img, bh_photo, bh_x, bh_y)
    d = ImageDraw.Draw(img)
    d.text((bh_x, bh_y - 18), "BH1750 (GY-302) — Sensor Cahaya", fill=ACCENT, font=get_font(12, True))

    # BH1750 connections
    bh_conn_x = bh_x + bh_photo.width + 15
    bh_conns = [
        ("VCC", "<-- 3.3V", ORANGE, bh_y + 20),
        ("GND", "<-- GND", WIRE_BLUE, bh_y + 45),
        ("SDA", "<-- GPIO8", ACCENT, bh_y + 70),
        ("SCL", "<-- GPIO9", ACCENT, bh_y + 95),
        ("ADDR", "(kosong)", DARK_GRAY, bh_y + 120),
    ]
    for pin, target, color, py in bh_conns:
        d.text((bh_conn_x, py), f"{pin} {target}", fill=color, font=get_font(10, True))

    # Wires BH1750
    mid_x = 320
    draw_wire(d, [(esp_right, pin_8_y), (mid_x, pin_8_y), (mid_x, bh_y+75), (bh_x, bh_y+75)], ACCENT, 3)
    draw_wire(d, [(esp_right, pin_9_y), (mid_x+15, pin_9_y), (mid_x+15, bh_y+100), (bh_x, bh_y+100)], ACCENT, 3)

    # ====== DHT11 REAL PHOTO ======
    dht_photo = load_real('dht11_module.jpg', max_h=120)
    dht_photo = add_shadow_border(dht_photo, border=2, shadow=4)
    dht_x, dht_y = 380, 230
    paste_real(img, dht_photo, dht_x, dht_y)
    d = ImageDraw.Draw(img)
    d.text((dht_x, dht_y - 18), "DHT11 Module — Sensor Suhu & Humidity", fill=GREEN, font=get_font(12, True))

    dht_conn_x = dht_x + dht_photo.width + 15
    dht_conns = [
        ("VCC", "<-- 3.3V", ORANGE, dht_y + 15),
        ("DATA", "<-- GPIO4", GREEN, dht_y + 45),
        ("GND", "<-- GND", WIRE_BLUE, dht_y + 75),
    ]
    for pin, target, color, py in dht_conns:
        d.text((dht_conn_x, py), f"{pin} {target}", fill=color, font=get_font(10, True))

    # Wire DHT11
    draw_wire(d, [(esp_right, pin_4_y), (mid_x-15, pin_4_y), (mid_x-15, dht_y+50), (dht_x, dht_y+50)], WIRE_GREEN, 3)

    # ====== RAIN SENSOR REAL PHOTO ======
    rain_photo = load_real('rain_module.jpg', max_h=100)
    rain_photo = add_shadow_border(rain_photo, border=2, shadow=4)
    rain_x, rain_y = 380, 395
    paste_real(img, rain_photo, rain_x, rain_y)
    d = ImageDraw.Draw(img)
    d.text((rain_x, rain_y - 18), "Rain Sensor + LM393 — Sensor Hujan", fill=YELLOW, font=get_font(12, True))

    rain_conn_x = rain_x + rain_photo.width + 15
    rain_conns = [
        ("VCC", "<-- 3.3V", ORANGE, rain_y + 10),
        ("GND", "<-- GND", WIRE_BLUE, rain_y + 35),
        ("DO", "<-- GPIO5", YELLOW, rain_y + 60),
        ("AO", "JANGAN!!!", RED, rain_y + 85),
    ]
    for pin, target, color, py in rain_conns:
        d.text((rain_conn_x, py), f"{pin} {target}", fill=color, font=get_font(10, True))

    # Wire Rain
    draw_wire(d, [(esp_right, pin_5_y), (mid_x-30, pin_5_y), (mid_x-30, rain_y+65), (rain_x, rain_y+65)], YELLOW, 3)

    # AO DANGER
    draw_warning_box(d, 380, h-80, 390, "Pin AO Rain Sensor = JANGAN DISAMBUNG!", "Output bisa 5V, RUSAK pin ADC ESP32 permanen!")

    # Legend
    draw_legend(d, 30, h-170, [
        (ORANGE, "3.3V (Power sensor)"),
        (WIRE_BLUE, "GND (Ground)"),
        (ACCENT, "I2C (SDA/SCL)"),
        (WIRE_GREEN, "Data DHT11"),
        (YELLOW, "Data Rain Sensor"),
        (RED, "BAHAYA! Jangan!"),
    ])

def img_servo_wiring(img, d, w, h):
    draw_title_bar(d, w, "Wiring Servo MG996R + Kapasitor 470uF", "POWER SERVO = LANGSUNG DARI ADAPTOR 5V!")

    # MG996R REAL PHOTO
    servo_photo = load_real('mg996r_real.jpg', max_h=220)
    servo_photo = add_shadow_border(servo_photo, border=3, shadow=6)
    paste_real(img, servo_photo, 40, 75)
    d = ImageDraw.Draw(img)
    d.text((40, 305), "Servo MG996R", fill=PURPLE, font=get_font(12, True))
    d.text((40, 322), "Torque: 11 kg.cm | 4.8-7.2V", fill=GRAY, font=get_font(9))

    # MG996R Pinout photo
    pinout = load_real('mg996r_pinout.png', max_h=180)
    pinout = add_shadow_border(pinout, border=2, shadow=4)
    paste_real(img, pinout, 320, 75)
    d = ImageDraw.Draw(img)
    d.text((320, 260), "Pinout MG996R (3 kabel warna)", fill=LIGHT_GRAY, font=get_font(10))

    # Wire descriptions
    wy = 300
    wires = [
        ("Kabel MERAH (VCC)", "+5V Power Rail (ADAPTOR!)", WIRE_RED, "Langsung dari adaptor, BUKAN ESP32!"),
        ("Kabel COKLAT (GND)", "GND Rail (ground bus)", WIRE_BLUE, "Ground bersama semua komponen"),
        ("Kabel ORANYE (Signal)", "GPIO13 ESP32", WIRE_ORANGE, "Sinyal PWM, kontrol posisi servo"),
    ]
    for i, (label, target, color, note) in enumerate(wires):
        wy_i = wy + 50 + i * 48
        d.rounded_rectangle([40, wy_i, w-40, wy_i+40], radius=6, fill=BG_CARD, outline=color, width=2)
        d.line([(50, wy_i+20), (80, wy_i+20)], fill=color, width=5)
        d.ellipse([76, wy_i+16, 84, wy_i+24], fill=color)
        d.text((92, wy_i+4), label, fill=color, font=get_font(11, True))
        draw_arrow(d, 300, wy_i+20, 340, wy_i+20, color, 2, 8)
        d.text((350, wy_i+4), target, fill=WHITE, font=get_font(11, True))
        d.text((350, wy_i+22), note, fill=GRAY, font=get_font(9))

    # Capacitor section
    cap_y = wy + 200
    d.rounded_rectangle([40, cap_y, w-40, cap_y+70], radius=8, fill=BG_CARD, outline=ACCENT, width=2)
    d.text((60, cap_y+8), "CAPACITOR 470uF 16V (Elektrolit)", fill=ACCENT, font=get_font(13, True))
    d.text((60, cap_y+28), "Kaki panjang (+) --> +5V Rail  |  Kaki pendek (-) --> GND Rail", fill=WHITE, font=get_font(11))
    d.text((60, cap_y+48), "Pasang DEKAT servo! Fungsi: peredam lonjakan arus saat servo start (2.5A stall)", fill=GRAY, font=get_font(9))

    draw_warning_box(d, 40, h-68, w-80, "SERVO POWER = ADAPTOR 5V! BUKAN dari pin ESP32!", "ESP32 max 500mA. Servo stall = 2500mA. Bisa RUSAK ESP32!")

def img_led_wiring(img, d, w, h):
    draw_title_bar(d, w, "Wiring LED Warning + Resistor 220 Ohm", "GPIO2 --> Resistor 220 Ohm --> LED Merah --> GND")

    # LED real photo
    led_photo = load_real('led_red.jpg', max_h=130)
    led_photo = add_shadow_border(led_photo, border=2, shadow=4)
    paste_real(img, led_photo, 40, 75)
    d = ImageDraw.Draw(img)
    d.text((40, 215), "LED Merah 5mm", fill=RED, font=get_font(11, True))
    d.text((40, 232), "Kaki panjang = + (Anoda)", fill=RED, font=get_font(9))
    d.text((40, 247), "Kaki pendek  = - (Katoda)", fill=BLUE, font=get_font(9))

    # Circuit diagram
    cy = 150
    d.text((300, 75), "RANGKAIAN SERI:", fill=WHITE, font=get_font(14, True))

    # GPIO2
    d.rounded_rectangle([280, cy-25, 370, cy+25], radius=8, fill=(18,65,48), outline=GREEN, width=2)
    d.text((295, cy-15), "ESP32", fill=GREEN, font=get_font(10, True))
    d.text((293, cy+2), "GPIO2", fill=PINK, font=get_font(12, True))

    draw_arrow(d, 370, cy, 400, cy, WIRE_GREEN, 3, 8)

    # Resistor
    d.rounded_rectangle([405, cy-20, 520, cy+20], radius=6, fill=(60,55,30), outline=(120,110,60), width=2)
    d.text((415, cy-15), "RESISTOR", fill=WHITE, font=get_font(9, True))
    d.text((420, cy+2), "220 Ohm", fill=YELLOW, font=get_font(10, True))

    draw_arrow(d, 520, cy, 550, cy, ORANGE, 3, 8)

    # LED symbol
    d.ellipse([555, cy-22, 605, cy+22], fill=(180,30,30), outline=(255,60,60), width=2)
    d.text((567, cy-10), "LED", fill=WHITE, font=get_font(10, True))
    d.text((565, cy+4), "Merah", fill=WHITE, font=get_font(8))

    draw_arrow(d, 605, cy, 635, cy, WIRE_BLUE, 3, 8)

    # GND
    d.rounded_rectangle([640, cy-18, 710, cy+18], radius=6, fill=(30,30,50), outline=BLUE, width=2)
    d.text((655, cy-8), "GND", fill=BLUE, font=get_font(12, True))

    # Flow arrow
    d.text((380, cy+35), "arus -->", fill=GREEN, font=get_font(9))

    # Info
    d.rounded_rectangle([280, 220, w-40, 330], radius=8, fill=BG_CARD, outline=DARK_GRAY)
    d.text((300, 228), "DETAIL:", fill=WHITE, font=get_font(12, True))
    details = [
        ("Resistor 220 Ohm:", "Tidak ada polaritas (boleh dibalik)", YELLOW),
        ("LED 5mm:", "Kaki PANJANG = Anoda (+), Kaki PENDEK = Katoda (-)", RED),
        ("Fungsi:", "Berkedip saat Anti-Oven warning aktif (tutup + panas > 33C)", ORANGE),
        ("Kenapa 220 Ohm?", "Membatasi arus ~15mA supaya LED tidak terbakar", ACCENT),
    ]
    for i, (label, desc, color) in enumerate(details):
        dy = 250 + i * 20
        d.text((300, dy), label, fill=color, font=get_font(10, True))
        d.text((460, dy), desc, fill=LIGHT_GRAY, font=get_font(10))

def img_checklist(img, d, w, h):
    draw_title_bar(d, w, "CHECKLIST SEBELUM POWER ON", "Cek SEMUA sebelum colok adaptor ke listrik!")

    items = [
        ("Semua GND terhubung ke SATU ground bus", GREEN, True),
        ("Tidak ada kabel 5V masuk langsung ke sensor (harus 3.3V!)", GREEN, True),
        ("Rain sensor pin AO TIDAK tersambung ke manapun", RED, True),
        ("Servo power dari ADAPTOR 5V, BUKAN dari pin ESP32", RED, True),
        ("Capacitor 470uF polaritas benar: (+) ke 5V, (-) ke GND", ORANGE, True),
        ("Tidak ada kabel/komponen yang saling menyentuh (short)", RED, True),
        ("ESP32 belum di-power (belum colok adaptor/USB)", WHITE, False),
        ("Kabel USB yang disiapkan = kabel DATA (bukan charge-only)", ORANGE, False),
        ("Semua jumper wire terpasang kuat, tidak longgar", WHITE, False),
    ]

    y = 68
    for text, color, critical in items:
        row_bg = (50, 25, 25) if critical else BG_CARD
        d.rounded_rectangle([35, y, w-35, y+42], radius=6, fill=row_bg, outline=(100,40,40) if critical else DARK_GRAY)
        d.rounded_rectangle([45, y+8, 72, y+35], radius=4, outline=color, width=2)
        d.line([(50, y+22), (55, y+28), (66, y+14)], fill=color, width=2)
        if critical:
            d.rounded_rectangle([80, y+8, 122, y+26], radius=3, fill=RED_DIM)
            d.text((86, y+9), "WAJIB", fill=RED, font=get_font(8, True))
            d.text((130, y+11), text, fill=color, font=get_font(11))
        else:
            d.text((80, y+11), text, fill=color, font=get_font(11))
        y += 48

    draw_warning_box(d, 35, h-65, w-70, "Kalau ada yang SALAH = komponen RUSAK PERMANEN!", "Double-check! Lebih baik lambat tapi aman.")

def img_serial_monitor(img, d, w, h):
    draw_title_bar(d, w, "Serial Monitor — Contoh Output Normal", "Tools > Serial Monitor | baud rate = 115200")

    mx, my = 30, 65
    mw, mh = w-60, h-120
    d.rounded_rectangle([mx, my, mx+mw, my+mh], radius=8, fill=BG_DARK, outline=(60,60,80), width=2)
    d.rectangle([mx, my, mx+mw, my+28], fill=(40,40,58))
    d.rounded_rectangle([mx, my, mx+mw, my+10], radius=8, fill=(40,40,58))
    d.text((mx+15, my+6), "Serial Monitor — COM3 — 115200 baud", fill=LIGHT_GRAY, font=get_font(10))

    mono = get_mono(11)
    y = my + 36
    lines = [
        ("============================================================", DARK_GRAY),
        ("  SMART GREENHOUSE IoT", ACCENT),
        ("  Board: ESP32-S3-N16R8 Soldered", ACCENT),
        ("  Firmware: All Phases Active", ACCENT),
        ("============================================================", DARK_GRAY),
        ("  Free heap  : 287456 bytes", GRAY),
        ("  Free PSRAM : 8388608 bytes", GRAY),
        ("============================================================", DARK_GRAY),
        ("", WHITE),
        ("[WIFI] Menghubungkan ke MyWiFi...", BLUE),
        ("[WIFI] Terhubung! IP: 192.168.1.47", GREEN),
        ("[WEATHER] Prediksi 6 jam: hujan=TIDAK (25%)", ACCENT),
        ("[SYSTEM] Inisialisasi selesai.", GREEN),
        ("", WHITE),
        ("------------------------------------------------------------", DARK_GRAY),
        ("  Suhu       : 28.5 C", WHITE),
        ("  Humidity   : 65.2 %", WHITE),
        ("  Cahaya     : 15420.0 lux", WHITE),
        ("  Hujan      : TIDAK", GREEN),
        ("  Atap       : 100% (180 derajat)", GREEN),
        ("  Mode       : AUTO", WHITE),
        ("  Overheat   : Normal", GREEN),
        ("  WiFi       : Terhubung", GREEN),
        ("  Forecast   : hujan=TIDAK (25%)", WHITE),
        ("  Free heap  : 245680 bytes", GRAY),
    ]
    for text, color in lines:
        if y > my + mh - 15: break
        d.text((mx+15, y), text, fill=color, font=mono)
        y += 16

    d.rounded_rectangle([30, h-48, w-30, h-8], radius=6, fill=BG_CARD, outline=DARK_GRAY)
    d.text((50, h-43), "Suhu/Humidity = -1.0 --> DHT11 error!", fill=RED, font=get_font(10, True))
    d.text((400, h-43), "Cahaya = -1.0 --> BH1750 error!", fill=RED, font=get_font(10, True))
    d.text((50, h-25), "Cek kabel sensor yang error tersebut.", fill=GRAY, font=get_font(10))

def img_supabase_setup(img, d, w, h):
    draw_title_bar(d, w, "Setup Supabase — Database Cloud Gratis", "supabase.com > Sign Up > New Project > Settings > API")

    steps = [
        ("1", "Buka supabase.com > Sign Up", "Bisa pakai akun GitHub", BLUE),
        ("2", "Klik 'New Project'", "Di dashboard utama", WHITE),
        ("3", "Nama: smart-greenhouse", "Nama project bebas", WHITE),
        ("4", "Region: Southeast Asia (Singapore)", "Terdekat dari Indonesia", ACCENT),
        ("5", "Set database password > CATAT!", "Simpan di tempat aman", RED),
        ("6", "Tunggu ~2 menit", "Server sedang provisioning", GRAY),
        ("7", "Settings > API > catat URL & Key", "2 hal penting (lihat bawah)", GREEN),
    ]

    y = 65
    for num, title, desc, color in steps:
        d.rounded_rectangle([35, y, w-35, y+42], radius=6, fill=BG_CARD, outline=DARK_GRAY)
        d.ellipse([45, y+6, 72, y+33], fill=ACCENT_DIM, outline=ACCENT, width=2)
        d.text((53, y+10), num, fill=WHITE, font=get_font(13, True))
        d.text((82, y+6), title, fill=color, font=get_font(12, True))
        d.text((82, y+24), desc, fill=GRAY, font=get_font(10))
        y += 48

    y += 5
    d.rounded_rectangle([35, y, w-35, y+100], radius=8, fill=BG_DARK, outline=GREEN, width=2)
    d.text((55, y+8), "CATAT 2 HAL INI dari Settings > API:", fill=GREEN, font=get_font(13, True))
    d.rounded_rectangle([55, y+32, w-55, y+55], radius=4, fill=(20,35,20), outline=GREEN_DIM)
    d.text((65, y+36), "Project URL:", fill=GRAY, font=get_font(10))
    d.text((170, y+36), "https://abcdefg.supabase.co", fill=GREEN, font=get_mono(11))
    d.rounded_rectangle([55, y+60, w-55, y+83], radius=4, fill=(20,35,20), outline=GREEN_DIM)
    d.text((65, y+64), "anon public:", fill=GRAY, font=get_font(10))
    d.text((170, y+64), "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...", fill=GREEN, font=get_mono(10))

    draw_warning_box(d, 55, y+88, w-110, "Pakai 'anon public' key! BUKAN 'service_role' key!", None)

def img_upload_firmware(img, d, w, h):
    draw_title_bar(d, w, "Upload Firmware ke ESP32-S3", "Klik Upload > Compile > Upload > Done!")

    steps = [
        ("1", "Buka smart_greenhouse.ino", "File > Open di Arduino IDE", WHITE, False),
        ("2", "Board = ESP32S3 Dev Module", "Tools > Board", WHITE, False),
        ("3", "Port = COM yang muncul", "Tools > Port", WHITE, False),
        ("4", "Klik UPLOAD", "Ikon panah kanan (Ctrl+U)", GREEN, False),
        ("5", "'Compiling sketch...'", "Tunggu 1-2 menit", GRAY, False),
        ("6", "'Uploading...'", "Progress bar berjalan", GRAY, False),
        ("7", "SELESAI!", "'Hard resetting via RTS pin...'", GREEN, True),
    ]

    y = 65
    for num, title, desc, color, done in steps:
        bg = (25, 50, 25) if done else BG_CARD
        d.rounded_rectangle([35, y, w-35, y+38], radius=6, fill=bg, outline=GREEN if done else DARK_GRAY)
        nc = GREEN if done else ACCENT
        d.ellipse([45, y+5, 68, y+28], fill=nc if done else ACCENT_DIM, outline=nc, width=2)
        d.text((51, y+8), num, fill=BG if done else WHITE, font=get_font(12, True))
        d.text((78, y+4), title, fill=color, font=get_font(11, True))
        d.text((78, y+21), desc, fill=GRAY, font=get_font(9))
        y += 44

    y += 8
    d.rounded_rectangle([35, y, w-35, y+105], radius=8, fill=(60,40,20), outline=ORANGE, width=2)
    d.text((55, y+8), "TROUBLESHOOT: Gagal 'Connecting........_____'", fill=ORANGE, font=get_font(12, True))
    boot_steps = [
        "1. Tahan tombol BOOT di board ESP32",
        "2. Sambil tahan BOOT, tekan RESET sekali",
        "3. Lepas RESET (BOOT masih ditahan)",
        "4. Klik Upload di Arduino IDE",
        "5. Setelah 'Uploading...' muncul, lepas BOOT",
    ]
    for i, s in enumerate(boot_steps):
        d.text((70, y+30+i*14), s, fill=LIGHT_GRAY, font=get_font(10))

def img_test_rain(img, d, w, h):
    draw_title_bar(d, w, "TEST: Hujan = Atap TUTUP (Rule 1: Local is King)", "Sensor hujan fisik = prioritas tertinggi, override segalanya")

    # Rain sensor real photo
    rain_photo = load_real('rain_board.png', max_h=120)
    rain_photo = add_shadow_border(rain_photo, border=2, shadow=4)
    paste_real(img, rain_photo, 40, 70)
    d = ImageDraw.Draw(img)
    d.text((40, 200), "Rain Sensor Module", fill=YELLOW, font=get_font(10, True))

    # Before
    d.rounded_rectangle([230, 70, 430, 200], radius=8, fill=(25,55,25), outline=GREEN, width=2)
    d.rounded_rectangle([230, 70, 430, 98], radius=8, fill=GREEN_DIM)
    d.rectangle([230, 88, 430, 98], fill=GREEN_DIM)
    d.text((245, 75), "SEBELUM", fill=WHITE, font=get_font(12, True))
    before = [("Hujan:", "TIDAK", GREEN), ("Atap:", "100% (180°)", GREEN), ("Mode:", "AUTO", WHITE)]
    for i, (k,v,c) in enumerate(before):
        d.text((245, 105+i*28), k, fill=GRAY, font=get_font(10))
        d.text((310, 105+i*28), v, fill=c, font=get_font(11, True))

    # Arrow + action
    draw_arrow(d, 430, 135, 490, 135, ORANGE, 3, 10)
    d.text((440, 110), "Tetes air\nke plate", fill=ORANGE, font=get_font(9))
    d.text((440, 150), "60 detik", fill=GRAY, font=get_font(8))

    # After
    d.rounded_rectangle([500, 70, 700, 200], radius=8, fill=(55,25,25), outline=RED, width=2)
    d.rounded_rectangle([500, 70, 700, 98], radius=8, fill=RED_DIM)
    d.rectangle([500, 88, 700, 98], fill=RED_DIM)
    d.text((515, 75), "SESUDAH (60 detik)", fill=WHITE, font=get_font(11, True))
    after = [("Hujan:", "YA", RED), ("Atap:", "0% (0°)", RED), ("Mode:", "AUTO", WHITE)]
    for i, (k,v,c) in enumerate(after):
        d.text((515, 105+i*28), k, fill=GRAY, font=get_font(10))
        d.text((580, 105+i*28), v, fill=c, font=get_font(11, True))

    # Explanation
    d.rounded_rectangle([40, 225, w-40, 300], radius=8, fill=BG_CARD, outline=ACCENT)
    d.text((60, 233), "Rule 1: Local is King", fill=ACCENT, font=get_font(14, True))
    d.text((60, 255), "Sensor hujan FISIK = prioritas TERTINGGI. Tidak bisa di-override oleh:", fill=LIGHT_GRAY, font=get_font(10))
    d.text((60, 275), "perintah manual, prediksi cuaca cerah, atau kondisi cahaya terang apapun.", fill=LIGHT_GRAY, font=get_font(10))

def img_final_system(img, d, w, h):
    draw_title_bar(d, w, "ARSITEKTUR SISTEM LENGKAP", "Semua komponen terhubung & bekerja otomatis")

    # Sensor photos (small)
    bh = load_real('bh1750.jpg', max_h=55)
    dht = load_real('dht11.jpg', max_h=55)
    rain = load_real('rain_module.jpg', max_h=55)
    servo = load_real('mg996r_module.jpg', max_h=55)
    esp = load_real('esp32s3.png', max_h=120)

    # Paste sensor photos
    paste_real(img, bh, 20, 75)
    paste_real(img, dht, 20, 145)
    paste_real(img, rain, 20, 215)
    d = ImageDraw.Draw(img)

    # Sensor labels
    sx = 20 + max(bh.width, dht.width, rain.width) + 5
    d.text((sx, 85), "BH1750", fill=ACCENT, font=get_font(10, True))
    d.text((sx, 100), "Cahaya", fill=GRAY, font=get_font(8))
    d.text((sx, 155), "DHT11", fill=GREEN, font=get_font(10, True))
    d.text((sx, 170), "Suhu/Hum", fill=GRAY, font=get_font(8))
    d.text((sx, 225), "Rain", fill=YELLOW, font=get_font(10, True))
    d.text((sx, 240), "Hujan", fill=GRAY, font=get_font(8))

    # Arrows to ESP
    for ay in [100, 170, 240]:
        draw_arrow(d, sx+60, ay, 240, ay, ACCENT if ay==100 else (GREEN if ay==170 else YELLOW), 2, 8)

    # ESP32 photo
    paste_real(img, esp, 245, 95)
    d = ImageDraw.Draw(img)
    d.text((250, 220), "ESP32-S3 N16R8", fill=GREEN, font=get_font(10, True))

    # Decision engine box
    eng_x = 250
    d.rounded_rectangle([eng_x, 240, eng_x+140, 330], radius=6, fill=BG_CARD, outline=ACCENT_DIM)
    d.text((eng_x+8, 245), "Decision Engine", fill=ACCENT, font=get_font(9, True))
    rules = [("R1:", "Local King", RED), ("R2:", "Hysteresis", ORANGE), ("R3:", "Failsafe", BLUE), ("R4:", "Anti-Oven", YELLOW)]
    for i, (rn, rd, rc) in enumerate(rules):
        d.text((eng_x+8, 262+i*16), rn, fill=rc, font=get_font(8, True))
        d.text((eng_x+30, 262+i*16), rd, fill=LIGHT_GRAY, font=get_font(8))

    # Servo + Roof
    paste_real(img, servo, 470, 80)
    d = ImageDraw.Draw(img)
    d.text((470, 80+servo.height+3), "MG996R", fill=PURPLE, font=get_font(9, True))
    draw_arrow(d, 245+esp.width, 130, 465, 110, PURPLE, 2, 8)

    d.rounded_rectangle([560, 80, 650, 140], radius=6, fill=(50,50,50), outline=GRAY)
    d.text((575, 90), "ATAP", fill=WHITE, font=get_font(12, True))
    d.text((570, 110), "Louvered", fill=GRAY, font=get_font(8))
    draw_arrow(d, 470+servo.width, 110, 555, 110, GRAY, 2, 8)

    # LED
    d.rounded_rectangle([470, 160, 570, 200], radius=6, fill=BG_CARD, outline=PINK)
    d.text((480, 166), "LED Warning", fill=PINK, font=get_font(10, True))
    d.text((480, 182), "Anti-Oven", fill=GRAY, font=get_font(8))
    draw_arrow(d, 245+esp.width, 175, 465, 180, PINK, 2, 8)

    # Cloud services
    cloud_y = 355
    d.rounded_rectangle([20, cloud_y, w-20, cloud_y+105], radius=10, fill=BG_CARD, outline=DARK_GRAY)
    d.text((40, cloud_y+5), "CLOUD SERVICES (via WiFi)", fill=BLUE, font=get_font(11, True))
    draw_arrow(d, 315, 330, 315, cloud_y, BLUE, 3, 10)
    d.text((325, 335), "WiFi", fill=BLUE, font=get_font(9, True))

    clouds = [
        (40, "Supabase", GREEN, "Database + Realtime"),
        (250, "Open-Meteo", ACCENT, "Prediksi Cuaca"),
        (460, "Vercel", PURPLE, "Web Dashboard"),
    ]
    for cx, name, color, desc in clouds:
        d.rounded_rectangle([cx, cloud_y+28, cx+190, cloud_y+90], radius=6, fill=BG_DARK, outline=color, width=2)
        d.text((cx+10, cloud_y+33), name, fill=color, font=get_font(13, True))
        d.text((cx+10, cloud_y+55), desc, fill=LIGHT_GRAY, font=get_font(9))
        d.text((cx+10, cloud_y+72), "sensor_logs, commands" if name=="Supabase" else ("6 jam ke depan" if name=="Open-Meteo" else "Monitor & Kontrol"), fill=GRAY, font=get_font(8))

# ============================================================
# Generate all images
# ============================================================
images = {
    'usb': (830, 480, img_usb_connect),
    'board': (850, 490, img_arduino_board),
    'libs': (780, 470, img_libraries),
    'breadboard': (800, 580, img_breadboard_esp),
    'power': (800, 520, img_power_wiring),
    'sensors': (800, 600, img_sensor_wiring),
    'servo': (750, 570, img_servo_wiring),
    'led': (750, 350, img_led_wiring),
    'checklist': (780, 520, img_checklist),
    'serial': (750, 520, img_serial_monitor),
    'supabase': (750, 520, img_supabase_setup),
    'upload': (730, 500, img_upload_firmware),
    'test_rain': (750, 320, img_test_rain),
    'final': (680, 480, img_final_system),
}

img_paths = {}
for name, (w, h, func) in images.items():
    img_paths[name] = make_img(name, w, h, func)
    print(f"  Generated: {name}.png")

# ============================================================
# WORD DOCUMENT (same structure as before)
# ============================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

step_num = [0]

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def step(title):
    step_num[0] += 1
    heading(f'Step {step_num[0]}: {title}', level=2)

def para(text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

def add_img(name, width=6.0):
    if name in img_paths:
        doc.add_picture(img_paths[name], width=Inches(width))

def warning(text):
    p = doc.add_paragraph()
    r = p.add_run(f'PERINGATAN: {text}')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

def tip(text):
    p = doc.add_paragraph()
    r = p.add_run(f'TIP: {text}')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x33, 0x99, 0x33)

# ======================== CONTENT ========================

doc.add_paragraph('')
h = doc.add_heading('PANDUAN PEMBUATAN\nSMART GREENHOUSE IoT', level=0)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
para('Step-by-Step dari Nol sampai Jadi', bold=True, size=14, color=(0x44,0x44,0x88))
para('ESP32-S3-N16R8 Soldered', bold=True, size=12, color=(0x44,0x44,0x88))
para('')
para('Panduan ini dirancang untuk orang awam.', size=11)
para('Ikuti step-by-step dari atas ke bawah.', size=11)
para('Jangan loncat step — setiap step bergantung pada step sebelumnya.', size=11)
para('')
para('13525004@std.stei.itb.ac.id | April 2026', size=10, color=(0x88,0x88,0x88))
doc.add_page_break()

heading('DAFTAR ISI', level=1)
for s in ['FASE 0: Persiapan Software (Step 1-3)', 'FASE 1: Rakit Hardware (Step 4-10)',
          'FASE 2: Upload & Test Firmware (Step 11-15)', 'FASE 3: Setup Supabase (Step 16-19)',
          'FASE 4: Test Sistem Lengkap (Step 20-24)', 'FASE 5: Web Dashboard (Step 25-27)']:
    para(s, size=11)
doc.add_page_break()

# FASE 0
heading('FASE 0: Persiapan Software', level=1)
para('Sebelum menyentuh hardware, siapkan software dulu.')

step('Download & Install Arduino IDE')
para('1. Buka browser, pergi ke arduino.cc/en/software')
para('2. Download Arduino IDE 2.x (pilih sesuai OS: Windows/Mac/Linux)')
para('3. Install seperti biasa (Next > Next > Finish)')
para('4. Buka Arduino IDE setelah install selesai')

step('Install ESP32 Board Support')
para('Arduino IDE default hanya kenal Arduino Uno. Kita perlu menambahkan ESP32.')
para('')
para('1. Buka Arduino IDE')
para('2. File > Preferences')
para('3. Di "Additional Board Manager URLs", paste:')
para('   https://espressif.github.io/arduino-esp32/package_esp32_index.json', bold=True)
para('4. OK > Tools > Board > Board Manager')
para('5. Cari "esp32" > "esp32 by Espressif Systems" > Install')
para('6. Tunggu selesai (5-10 menit)')

step('Install Library yang Dibutuhkan')
para('Library = kode tambahan supaya ESP32 bisa berkomunikasi dengan sensor.')
add_img('libs')
para('')
para('Tools > Manage Libraries, cari & install:')
para('   a. "BH1750" by Christopher Laws')
para('   b. "DHT sensor library" by Adafruit')
para('   c. "Adafruit Unified Sensor" by Adafruit')
para('   d. "ESP32Servo" by Kevin Harrington')
para('   e. "ArduinoJson" by Benoit Blanchon')
doc.add_page_break()

# FASE 1
heading('FASE 1: Rakit Hardware', level=1)
para('Rakit rangkaian elektronik di breadboard.')
warning('Jangan colok adaptor/USB sampai semua wiring selesai dan dicek!')

step('Pasang ESP32 di Breadboard')
add_img('breadboard')
para('')
para('1. Ambil breadboard 830 point')
para('2. Tancapkan ESP32-S3 di TENGAH breadboard')
para('3. Kedua sisi pin harus bisa diakses untuk jumper wire')
tip('Tekan kuat tapi hati-hati. Semua pin harus masuk.')

step('Sambung Power Supply')
add_img('power')
para('')
para('1. Adaptor 5V 3A > jack DC female screw terminal (belum colok!)')
para('2. Terminal (+) > power rail (+) breadboard = +5V', bold=True)
para('3. Terminal (-) > ground rail (-) breadboard = GND', bold=True)
para('4. Power rail (+) > ESP32 pin VIN')
para('5. Ground rail (-) > ESP32 pin GND')
para('6. ESP32 pin 3.3V > baris baru breadboard (jadi 3.3V bus sensor)')
warning('VIN ESP32 = 5V. Sensor = 3.3V. Jangan tukar!')

step('Sambung Sensor Cahaya (BH1750)')
add_img('sensors')
para('')
para('BH1750 pakai I2C (SDA + SCL):')
para('  VCC  -->  3.3V bus', bold=True)
para('  GND  -->  GND bus', bold=True)
para('  SDA  -->  GPIO8', bold=True)
para('  SCL  -->  GPIO9', bold=True)
para('  ADDR -->  kosong')

step('Sambung Sensor Suhu (DHT11)')
para('DHT11 module 3-pin:')
para('  VCC  -->  3.3V bus', bold=True)
para('  GND  -->  GND bus', bold=True)
para('  DATA -->  GPIO4', bold=True)
tip('Beli yang MODULE 3-pin (ada PCB), bukan komponen telanjang 4-pin.')

step('Sambung Sensor Hujan (Rain Sensor)')
para('Dari control board LM393:')
para('  VCC  -->  3.3V bus', bold=True)
para('  GND  -->  GND bus', bold=True)
para('  DO   -->  GPIO5', bold=True)
para('  AO   -->  JANGAN DISAMBUNG!', bold=True, color=(0xCC,0x33,0x33))
warning('Pin AO bisa output 5V. ESP32 ADC max 3.3V. Disambung = RUSAK PERMANEN.')

step('Sambung Servo Motor (MG996R)')
add_img('servo')
para('')
para('  MERAH (VCC)   -->  +5V Rail (ADAPTOR!)', bold=True)
para('  COKLAT (GND)  -->  GND bus', bold=True)
para('  ORANYE (Signal) -->  GPIO13', bold=True)
para('')
para('Capacitor 470uF dekat servo:')
para('  Kaki panjang (+) --> +5V Rail')
para('  Kaki pendek (-) --> GND Rail')
warning('Servo power WAJIB dari adaptor. BUKAN ESP32! Servo butuh 2.5A, ESP32 max 500mA.')

step('Sambung LED Warning + Resistor')
add_img('led')
para('')
para('GPIO2 > Resistor 220 Ohm > LED Merah > GND')
para('Resistor: tidak ada polaritas. LED: kaki panjang = (+).')
doc.add_page_break()

# FASE 2
heading('FASE 2: Upload & Test Firmware', level=1)

step('Checklist Sebelum Power On')
add_img('checklist')
para('')
para('Cek SEMUA item sebelum colok adaptor!', bold=True)
para('Setelah yakin:')
para('1. Colok adaptor > cek LED power ESP32 menyala')
para('2. Tunggu 10 detik > sentuh komponen')
para('3. Kalau ada yang PANAS > CABUT SEGERA')

step('Colok ESP32 ke Laptop via USB')
add_img('usb')
para('')
para('1. Kabel USB-C DATA (bukan charge-only!)')
para('2. Colok ke ESP32 dan laptop')
para('3. Device Manager harus muncul COM port baru')
para('')
para('Tidak muncul? Ganti kabel USB atau install driver CH340/CP2102.')

step('Konfigurasi Arduino IDE')
add_img('board')
para('')
para('  Board: ESP32S3 Dev Module', bold=True)
para('  PSRAM: OPI PSRAM', bold=True, color=(0xCC,0x33,0x33))
para('  Flash Size: 16MB (128Mb)', bold=True)
para('  USB CDC On Boot: Enabled', bold=True, color=(0xCC,0x33,0x33))
para('  Upload Speed: 921600', bold=True)
para('  Port: COM yang muncul', bold=True)
warning('PSRAM = OPI & USB CDC = Enabled WAJIB BENAR!')

step('Siapkan File Firmware')
para('1. Clone/download folder firmware/ dari repo')
para('2. Buka smart_greenhouse.ino di Arduino IDE')
para('3. Copy credentials.h.example > credentials.h')
para('4. Edit credentials.h:')
para('   #define WIFI_SSID     "nama_wifi"', bold=True)
para('   #define WIFI_PASSWORD "password_wifi"', bold=True)
para('   #define SUPABASE_URL  "https://xxx.supabase.co"', bold=True)
para('   #define SUPABASE_KEY  "eyJ..."', bold=True)
tip('Isi WiFi dulu. Supabase nanti. Sistem tetap jalan offline.')

step('Upload Firmware ke ESP32')
add_img('upload')
para('')
para('Klik Upload > tunggu compile > tunggu upload.')
para('Kalau gagal "Connecting...": tahan BOOT > tekan RESET > lepas > upload > lepas BOOT')

step('Buka Serial Monitor')
add_img('serial')
para('')
para('Tools > Serial Monitor, baud rate 115200.')
para('Tekan RESET di ESP32. Harus muncul data sensor.')
para('Kalau ada -1 = sensor error, cek kabelnya.', color=(0xCC,0x33,0x33))
doc.add_page_break()

# FASE 3
heading('FASE 3: Setup Supabase', level=1)

step('Buat Akun & Project Supabase')
add_img('supabase')
para('')
para('supabase.com > Sign Up > New Project > smart-greenhouse')
para('Region: Southeast Asia (Singapore)')

step('Catat URL dan API Key')
para('Settings > API:')
para('  Project URL: https://xxxxx.supabase.co', bold=True)
para('  anon public key: eyJ...', bold=True)
para('Update credentials.h > upload ulang firmware.')
warning('Pakai "anon public" BUKAN "service_role" key!')

step('Buat Tabel Database')
para('SQL Editor > New Query > paste & Run:')
sql = """CREATE TABLE sensor_logs (
  id BIGSERIAL PRIMARY KEY, temperature FLOAT, humidity FLOAT,
  lux FLOAT, is_raining BOOLEAN DEFAULT FALSE,
  vent_position INTEGER DEFAULT 0,
  roof_angle INTEGER DEFAULT 0,
  mode TEXT DEFAULT 'auto', overheating BOOLEAN DEFAULT FALSE,
  created_at TIMESTAMPTZ DEFAULT NOW()
);
CREATE TABLE commands (
  id BIGSERIAL PRIMARY KEY,
  vent_position INTEGER CHECK (vent_position BETWEEN 0 AND 100),
  mode TEXT DEFAULT 'auto',
  executed BOOLEAN DEFAULT FALSE,
  created_at TIMESTAMPTZ DEFAULT NOW(), executed_at TIMESTAMPTZ
);
ALTER TABLE sensor_logs ENABLE ROW LEVEL SECURITY;
ALTER TABLE commands ENABLE ROW LEVEL SECURITY;
CREATE POLICY "read" ON sensor_logs FOR SELECT USING (true);
CREATE POLICY "insert" ON sensor_logs FOR INSERT WITH CHECK (true);
CREATE POLICY "all" ON commands FOR ALL USING (true) WITH CHECK (true);"""
p = doc.add_paragraph()
r = p.add_run(sql)
r.font.name = 'Consolas'
r.font.size = Pt(8)

step('Verifikasi Data Masuk')
para('Tunggu 60 detik > Table Editor > sensor_logs > harus ada data baru!')
doc.add_page_break()

# FASE 4
heading('FASE 4: Test Sistem Lengkap', level=1)

step('Test: Sensor Baca Normal')
para('Serial Monitor tiap 5 detik: suhu (20-35C), humidity (40-90%), cahaya (0-65000 lux).')

step('Test: Hujan = Atap Tutup (Rule 1)')
add_img('test_rain')
para('')
para('Teteskan air ke rain plate > "Hujan: YA" > tunggu 60 detik > servo ke posisi 0% (0°).')
para('Rule 1: Local is King. Tidak bisa di-override.')

step('Test: Cerah = Atap Buka')
para('Keringkan rain plate > arahkan senter ke BH1750 > tunggu 60 detik > servo ke posisi 100% (180°).')

step('Test: Hysteresis')
para('Tetes air > tunggu 30 detik > keringkan > servo TIDAK BOLEH bergerak!')

step('Test: Anti-Oven (Rule 4)')
para('Hujan aktif + suhu > 33C: LED berkedip, atap tetap tutup.')
doc.add_page_break()

# FASE 5
heading('FASE 5: Web Dashboard (Opsional)', level=1)

step('Setup Next.js')
para('npx create-next-app@latest greenhouse-dashboard', bold=True)
para('npm install @supabase/supabase-js', bold=True)
para('.env.local: NEXT_PUBLIC_SUPABASE_URL & NEXT_PUBLIC_SUPABASE_ANON_KEY')

step('Deploy ke Vercel')
para('Push ke GitHub > vercel.com > Import > Add env vars > Deploy')

step('Test Dashboard')
para('Data sensor muncul, slider posisi ventilasi 0-100% berfungsi, tombol AUTO/MANUAL aktif.')
doc.add_page_break()

# FINAL
heading('SISTEM LENGKAP', level=1)
add_img('final')
para('')
para('Selamat! Smart Greenhouse IoT sudah berjalan.', bold=True, size=12)
para('Sensor tiap 5s | Supabase tiap 60s | Command tiap 10s | Weather tiap 30min | Servo 0-100% | Offline OK')
para('')
para('13525004@std.stei.itb.ac.id | Smart Greenhouse IoT | April 2026', size=9, color=(0x88,0x88,0x88))

# SAVE
output_path = os.path.join(OUT_DIR, 'Panduan_Pembuatan_Smart_Greenhouse.docx')
doc.save(output_path)
print(f'\nDocument saved: {output_path}')
