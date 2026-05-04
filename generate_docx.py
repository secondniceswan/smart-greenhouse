from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Helpers
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

def detail_table(rows):
    table = doc.add_table(rows=len(rows), cols=2, style='Light Grid Accent 1')
    for i, (k, v) in enumerate(rows):
        ck = table.rows[i].cells[0]
        cv = table.rows[i].cells[1]
        pk = ck.paragraphs[0]
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        rk.font.name = 'Calibri'
        pv = cv.paragraphs[0]
        rv = pv.add_run(v)
        rv.font.size = Pt(10)
        rv.font.name = 'Calibri'
    doc.add_paragraph('')

def summary_table(headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Medium Shading 1 Accent 1')
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    doc.add_paragraph('')
    return table

# ========== COVER ==========
doc.add_paragraph('')
doc.add_paragraph('')
heading('DAFTAR KOMPONEN LENGKAP', level=0)
para('Smart Greenhouse IoT', bold=True, size=16, color=(0x44, 0x44, 0x88), align='center')
para('ESP32-S3-N16R8 Soldered', bold=True, size=14, color=(0x44, 0x44, 0x88), align='center')
para('Dokumen Teknis | Versi 1.0 | April 2026', italic=True, size=10, color=(0x88, 0x88, 0x88), align='center')
para('13525004@std.stei.itb.ac.id', italic=True, size=10, color=(0x88, 0x88, 0x88), align='center')
doc.add_page_break()

# ========== 1. MICROCONTROLLER ==========
heading('1. Microcontroller', level=1)
detail_table([
    ('Nama', 'ESP32-S3-N16R8 Soldered Development Board'),
    ('Chip', 'ESP32-S3 (Xtensa LX7 Dual-Core 240MHz)'),
    ('Flash', '16MB'),
    ('PSRAM', '8MB (Octal SPI)'),
    ('SRAM', '512KB'),
    ('WiFi', '802.11 b/g/n 2.4GHz'),
    ('Bluetooth', 'BLE 5.0'),
    ('USB', 'USB-OTG native (tanpa chip UART eksternal)'),
    ('GPIO', '45 pin'),
    ('Operating Voltage', '3.3V logic, 5V input via VIN/USB'),
    ('Jumlah', '1 buah'),
    ('Keyword Shopee', 'ESP32-S3-N16R8 development board'),
    ('Est. Harga', 'Rp 90.000 - 140.000'),
    ('Catatan', 'Pastikan yang N16R8 (bukan N8R2 atau N4R2). Cek tulisan di chip module. Board ESP32-S3 DevKitC generic juga kompatibel.'),
])

# ========== 2. SENSOR CAHAYA ==========
heading('2. Sensor Cahaya', level=1)
detail_table([
    ('Nama', 'Light Intensity Sensor Module'),
    ('Module', 'GY-302'),
    ('Chip', 'BH1750FVI (ROHM)'),
    ('Interface', 'I2C (SDA + SCL)'),
    ('I2C Address', '0x23 (default, ADDR pin floating)'),
    ('Range', '1 - 65.535 lux'),
    ('Resolusi', '1 lux'),
    ('Operating Voltage', '3.3V - 5V'),
    ('Arus', '0.12mA'),
    ('Pin Module', 'VCC, GND, SDA, SCL, ADDR'),
    ('Jumlah', '1 buah'),
    ('Keyword Shopee', 'BH1750 GY-302 sensor cahaya'),
    ('Est. Harga', 'Rp 8.000 - 15.000'),
    ('Catatan', 'Beli yang module GY-302 (PCB biru/ungu kecil). Sudah ada regulator onboard, aman di 3.3V.'),
])

# ========== 3. SENSOR SUHU ==========
heading('3. Sensor Suhu & Kelembapan', level=1)
detail_table([
    ('Nama', 'Temperature & Humidity Sensor Module'),
    ('Module', 'DHT11 Module (3-pin, PCB mount)'),
    ('Chip', 'DHT11 (Aosong)'),
    ('Interface', 'Single-wire digital (1 data pin)'),
    ('Range Suhu', '0 C - 50 C (akurasi +/- 2 C)'),
    ('Range Kelembapan', '20% - 90% RH (akurasi +/- 5%)'),
    ('Sampling Rate', '1 reading per detik (max)'),
    ('Operating Voltage', '3.3V - 5.5V'),
    ('Arus', '0.5 - 2.5mA'),
    ('Pin Module', 'VCC, DATA, GND (3 pin)'),
    ('Jumlah', '1 buah'),
    ('Keyword Shopee', 'DHT11 module 3 pin'),
    ('Est. Harga', 'Rp 8.000 - 15.000'),
    ('Catatan', 'WAJIB beli yang module 3-pin (sudah ada pull-up resistor 10K ohm di PCB). JANGAN beli yang komponen telanjang 4-pin.'),
])

# ========== 4. SENSOR HUJAN ==========
heading('4. Sensor Hujan', level=1)
detail_table([
    ('Nama', 'Raindrop Detection Sensor Module'),
    ('Module', 'MH-RD Rain Sensor + LM393 Comparator Board'),
    ('Terdiri dari', '2 bagian: (1) Rain detection plate (PCB terbuka), (2) Control board (LM393)'),
    ('Chip Comparator', 'LM393 (Texas Instruments)'),
    ('Output Digital (DO)', 'HIGH = tidak hujan, LOW = hujan (default, adjustable via potensiometer)'),
    ('Output Analog (AO)', '0V - VCC (TIDAK DIPAKAI - bahaya untuk ESP32 ADC yang max 3.3V)'),
    ('Sensitivity', 'Adjustable via potensiometer onboard'),
    ('Operating Voltage', '3.3V - 5V'),
    ('Arus', '~15mA'),
    ('Pin Module', 'VCC, GND, DO, AO'),
    ('Jumlah', '1 set (plate + control board + kabel penghubung)'),
    ('Keyword Shopee', 'raindrop sensor module LM393'),
    ('Est. Harga', 'Rp 10.000 - 18.000'),
    ('Catatan', 'Yang dipakai HANYA pin DO (digital output). Pin AO JANGAN disambung ke ESP32 - output bisa mencapai 5V, melebihi batas ADC 3.3V ESP32 dan bisa merusak pin.'),
])

# ========== 5. SERVO ==========
heading('5. Servo Motor (Aktuator Atap Louver)', level=1)
detail_table([
    ('Nama', 'High Torque Metal Gear Servo Motor'),
    ('Model', 'MG996R'),
    ('Merk Umum', 'TowerPro MG996R (atau klon kompatibel)'),
    ('Jenis Gear', 'Metal gear (bukan plastik)'),
    ('Torsi', '10 kg.cm (4.8V) / 11 kg.cm (6V)'),
    ('Sudut Rotasi', '0 - 180 derajat'),
    ('Operating Voltage', '4.8V - 7.2V'),
    ('Arus Jalan (running)', '500 - 900mA'),
    ('Arus Stall (macet)', '~2.5A'),
    ('Sinyal', 'PWM, menerima logic 3.3V dari ESP32'),
    ('Kecepatan', '0.17 detik / 60 derajat (pada 4.8V)'),
    ('Kabel', '3 wire: Merah (VCC), Coklat (GND), Oranye (Signal)'),
    ('Konektor', 'JR/Futaba 3-pin standard'),
    ('Jumlah', '1 buah'),
    ('Keyword Shopee', 'MG996R servo motor metal gear'),
    ('Est. Harga', 'Rp 35.000 - 55.000'),
    ('Catatan', 'Power WAJIB dari adaptor 5V langsung, BUKAN dari pin ESP32. Pin 5V ESP32 max ~500mA, servo butuh sampai 2.5A. Pastikan beli MG996R (metal gear), bukan SG90 (plastik gear, torsi kecil - tidak kuat buka atap).'),
])

# ========== 6. ADAPTOR ==========
heading('6. Adaptor Power Supply (AC to DC)', level=1)
detail_table([
    ('Nama', 'AC-DC Switching Power Adapter'),
    ('Input (colokan listrik)', 'AC 100-240V 50/60Hz (universal, cocok PLN Indonesia 220V)'),
    ('Output (ke rangkaian)', 'DC 5V 3A (3000mA / 15 Watt)'),
    ('Konektor Output', 'Jack DC barrel 5.5mm x 2.1mm (center positive)'),
    ('Regulasi', 'Regulated (output stabil 5V)'),
    ('Proteksi', 'Over-current, over-voltage, short-circuit protection'),
    ('Bentuk Fisik', 'Seperti charger laptop kecil - colokan 2 pin ke stopkontak, keluar kabel dengan ujung jack DC bulat'),
    ('Jumlah', '1 buah'),
    ('Keyword Shopee', 'adaptor DC 5V 3A jack 5.5mm'),
    ('Est. Harga', 'Rp 20.000 - 35.000'),
    ('Catatan', 'WAJIB 5V 3A minimum. Budget arus: ESP32 ~250mA + Servo ~2500mA peak = 2750mA, jadi 3A minimum. Pastikan regulated (output stabil), bukan adaptor abal-abal yang output bisa naik ke 6-7V.'),
])

# ========== 7. JACK DC ==========
heading('7. Jack DC Female to Screw Terminal', level=1)
detail_table([
    ('Nama', 'DC Power Jack Female Adapter to Screw Terminal'),
    ('Ukuran Jack', '5.5mm x 2.1mm (match dengan adaptor di atas)'),
    ('Output', '2 screw terminal (+ dan -)'),
    ('Fungsi', 'Mengkonversi jack DC bulat dari adaptor menjadi 2 kabel yang bisa masuk breadboard via jumper wire'),
    ('Cara Pakai', 'Colok jack DC adaptor ke lubang female, lalu sambung kabel jumper ke 2 screw terminal (+ dan -)'),
    ('Jumlah', '1 buah'),
    ('Keyword Shopee', 'jack DC female 5.5mm screw terminal'),
    ('Est. Harga', 'Rp 3.000 - 5.000'),
    ('Catatan', 'Tanpa komponen ini, Anda harus memotong kabel adaptor (tidak direkomendasikan). Ini solusi bersih dan aman.'),
])

# ========== 8. KABEL USB ==========
heading('8. Kabel USB Data', level=1)
detail_table([
    ('Nama', 'USB Data Cable'),
    ('Tipe Konektor', 'USB-C to USB-A (board ESP32-S3 Soldered biasanya USB-C)'),
    ('Panjang', '1 meter (cukup untuk dari laptop ke breadboard)'),
    ('Fungsi', '(1) Upload firmware dari PC/laptop ke ESP32, (2) Serial Monitor untuk debugging'),
    ('Jumlah', '1 buah'),
    ('Keyword Shopee', 'kabel USB C data 1 meter'),
    ('Est. Harga', 'Rp 8.000 - 15.000'),
    ('Catatan', 'KRITIS: Harus kabel DATA, bukan charge-only. Banyak kabel HP murah yang charge-only (tidak ada jalur data di dalamnya). Cara test: colok ESP32 ke PC, buka Device Manager, harus muncul COM port baru. Kalau tidak muncul = kabel charge-only, ganti kabel.'),
])

# ========== 9. JUMPER M-M ==========
heading('9. Kabel Jumper Male-to-Male', level=1)
detail_table([
    ('Nama', 'Dupont Jumper Wire Male-to-Male'),
    ('Panjang', '20cm'),
    ('Isi', '40 pcs per set (berbagai warna: merah, hitam, kuning, hijau, biru, dll)'),
    ('Fungsi', 'Menghubungkan antar lubang di breadboard, atau dari pin ESP32 ke breadboard power rail'),
    ('Jumlah', '1 set (40 pcs)'),
    ('Keyword Shopee', 'kabel jumper male male 20cm 40pcs'),
    ('Est. Harga', 'Rp 8.000 - 12.000'),
    ('Catatan', 'Kedua ujung berupa pin tajam (male). Gunakan warna berbeda untuk membedakan jalur: merah = 5V, hitam = GND, warna lain = sinyal.'),
])

# ========== 10. JUMPER M-F ==========
heading('10. Kabel Jumper Male-to-Female', level=1)
detail_table([
    ('Nama', 'Dupont Jumper Wire Male-to-Female'),
    ('Panjang', '20cm'),
    ('Isi', '40 pcs per set (berbagai warna)'),
    ('Fungsi', 'Menghubungkan dari sensor module yang memiliki pin header (male) ke lubang breadboard'),
    ('Jumlah', '1 set (40 pcs)'),
    ('Keyword Shopee', 'kabel jumper male female 20cm 40pcs'),
    ('Est. Harga', 'Rp 8.000 - 12.000'),
    ('Catatan', 'Satu ujung pin (male), satu ujung lubang (female). Dipakai untuk menyambung sensor seperti BH1750, DHT11, rain module ke breadboard.'),
])

# ========== 11. BREADBOARD ==========
heading('11. Breadboard', level=1)
detail_table([
    ('Nama', 'Solderless Breadboard'),
    ('Ukuran', 'Full-size, 830 tie points'),
    ('Dimensi', '~165mm x 55mm'),
    ('Bus Power', '2 rail (+ dan -) di setiap sisi, total 4 rail'),
    ('Fungsi', 'Tempat merakit rangkaian elektronik tanpa perlu solder. Lubang-lubang di dalamnya terhubung secara internal sesuai baris/kolom.'),
    ('Jumlah', '1 buah'),
    ('Keyword Shopee', 'breadboard 830 point'),
    ('Est. Harga', 'Rp 12.000 - 20.000'),
    ('Catatan', 'Beli yang 830 tie points (full-size). Jangan yang 400 (half-size) - tidak cukup lebar untuk ESP32 DevKit + semua sensor.'),
])

# ========== 12. CAPACITOR ==========
heading('12. Capacitor Elektrolit', level=1)
detail_table([
    ('Nama', 'Electrolytic Capacitor'),
    ('Nilai', '470 uF (microfarad)'),
    ('Tegangan Max', '16V (lebih dari cukup untuk rangkaian 5V)'),
    ('Polaritas', 'POLARIZED - kaki panjang = positif (+), kaki pendek = negatif (-), ada strip putih di body = sisi negatif (-)'),
    ('Fungsi', 'Menyimpan energi sesaat untuk meng-buffer lonjakan arus ketika servo mulai bergerak. Tanpa capacitor, lonjakan arus servo bisa menyebabkan voltage drop yang membuat ESP32 restart.'),
    ('Posisi Pemasangan', 'Dipasang sedekat mungkin dengan sambungan kabel power servo (VCC dan GND servo)'),
    ('Jumlah', '1 buah'),
    ('Keyword Shopee', 'capacitor elektrolit 470uF 16V'),
    ('Est. Harga', 'Rp 2.000 - 5.000'),
    ('Catatan', 'PERINGATAN: Polaritas JANGAN terbalik. Kaki panjang (+) ke jalur +5V, kaki pendek (-) ke jalur GND. Jika terbalik, capacitor bisa meledak kecil dan mengeluarkan gas/cairan.'),
])

# ========== 13. LED ==========
heading('13. LED Merah 5mm', level=1)
detail_table([
    ('Nama', 'LED (Light Emitting Diode)'),
    ('Warna', 'Merah'),
    ('Ukuran', '5mm diameter (standard through-hole)'),
    ('Forward Voltage', '~2.0V'),
    ('Forward Current', '20mA (max)'),
    ('Fungsi', 'Indikator warning saat terjadi anti-oven effect (atap tertutup karena hujan tapi suhu di dalam greenhouse lebih dari 33 C). LED akan berkedip sebagai tanda peringatan.'),
    ('Identifikasi Kaki', 'Kaki panjang = Anoda (+), kaki pendek = Katoda (-). Jika sudah dipotong, lihat bagian dalam LED: sisi kecil = anoda, sisi lebar = katoda.'),
    ('Jumlah', '2 buah (1 dipakai, 1 cadangan jika rusak)'),
    ('Keyword Shopee', 'LED 5mm merah'),
    ('Est. Harga', 'Rp 500 - 1.000 per pcs'),
    ('Catatan', 'LED WAJIB dipasang bersama resistor 220 ohm (komponen nomor 14). Tanpa resistor, arus terlalu besar dan LED langsung terbakar.'),
])

# ========== 14. RESISTOR ==========
heading('14. Resistor 220 Ohm', level=1)
detail_table([
    ('Nama', 'Resistor Carbon Film'),
    ('Nilai', '220 Ohm'),
    ('Kode Warna', 'Merah - Merah - Coklat - Emas (4 gelang)'),
    ('Toleransi', '+/- 5% (ditandai gelang emas di ujung)'),
    ('Daya', '1/4 Watt (standard, ukuran kecil)'),
    ('Fungsi', 'Membatasi arus yang mengalir ke LED. Perhitungan: (3.3V - 2.0V) / 220 Ohm = ~6mA. Ini aman untuk LED (max 20mA) dan untuk pin GPIO ESP32 (max 40mA).'),
    ('Cara Pasang', 'Dipasang seri antara pin GPIO ESP32 dan kaki anoda (+) LED. Resistor tidak punya polaritas, boleh dibalik.'),
    ('Jumlah', '2 buah (1 dipakai, 1 cadangan)'),
    ('Keyword Shopee', 'resistor 220 ohm 1/4W'),
    ('Est. Harga', 'Rp 500 - 1.000 per pack'),
])

# ========== 15. TERMINAL BLOCK ==========
heading('15. Terminal Block (Opsional)', level=1)
detail_table([
    ('Nama', 'PCB Screw Terminal Block'),
    ('Pin', '2-pin'),
    ('Pitch', '5.08mm (standard)'),
    ('Fungsi', 'Sambungan power yang lebih rapi dan kuat dari jumper wire biasa. Cocok untuk jalur arus besar seperti power servo. Kabel dimasukkan lalu dikunci dengan sekrup kecil.'),
    ('Jumlah', '2 buah'),
    ('Keyword Shopee', 'terminal block 2 pin PCB 5mm'),
    ('Est. Harga', 'Rp 1.500 - 2.500 per pcs'),
    ('Catatan', 'Opsional tapi direkomendasikan. Jumper wire biasa sudah cukup untuk prototyping, tapi terminal block lebih aman untuk jalur arus tinggi.'),
])

# ========== RINGKASAN ==========
doc.add_page_break()
heading('RINGKASAN TOTAL KOMPONEN', level=1)

summary_table(
    ['No', 'Komponen', 'Qty', 'Est. Harga (IDR)'],
    [
        ['1', 'ESP32-S3-N16R8 Soldered', '1', '90.000 - 140.000'],
        ['2', 'BH1750 GY-302 (Sensor Cahaya)', '1', '8.000 - 15.000'],
        ['3', 'DHT11 Module 3-pin (Suhu & Humidity)', '1', '8.000 - 15.000'],
        ['4', 'Rain Sensor MH-RD + LM393', '1 set', '10.000 - 18.000'],
        ['5', 'Servo MG996R (Metal Gear)', '1', '35.000 - 55.000'],
        ['6', 'Adaptor DC 5V 3A (AC 220V to DC 5V)', '1', '20.000 - 35.000'],
        ['7', 'Jack DC Female Screw Terminal', '1', '3.000 - 5.000'],
        ['8', 'Kabel USB-C Data (bukan charge-only)', '1', '8.000 - 15.000'],
        ['9', 'Jumper Male-Male 20cm (40pcs)', '1 set', '8.000 - 12.000'],
        ['10', 'Jumper Male-Female 20cm (40pcs)', '1 set', '8.000 - 12.000'],
        ['11', 'Breadboard 830 point (full-size)', '1', '12.000 - 20.000'],
        ['12', 'Capacitor 470uF 16V', '1', '2.000 - 5.000'],
        ['13', 'LED Merah 5mm', '2', '1.000 - 2.000'],
        ['14', 'Resistor 220 Ohm 1/4W', '2', '500 - 1.000'],
        ['15', 'Terminal Block 2-pin (opsional)', '2', '3.000 - 5.000'],
    ]
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TOTAL ESTIMASI: Rp 216.500 - Rp 355.000')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1A, 0x6B, 0x1A)

# ========== CATATAN PENTING ==========
doc.add_paragraph('')
heading('CATATAN PENTING', level=1)

notes = [
    ('Adaptor', 'Input AC 220V (colok ke stopkontak PLN biasa) menghasilkan output DC 5V 3A. Bentuknya seperti charger laptop kecil dengan ujung jack DC bulat.'),
    ('Servo MG996R', 'Power WAJIB langsung dari adaptor 5V, BUKAN dari pin ESP32. Pin 5V ESP32 max ~500mA, sedangkan servo butuh sampai 2.5A saat stall.'),
    ('Rain Sensor', 'Pakai HANYA pin DO (digital). Pin AO (analog) JANGAN PERNAH disambung ke ESP32. Output AO bisa mencapai 5V dan akan merusak pin ADC ESP32 yang max 3.3V.'),
    ('Capacitor 470uF', 'Polaritas jangan terbalik. Kaki panjang (+) ke jalur +5V, kaki pendek (-) ke jalur GND. Terbalik bisa meledak kecil.'),
    ('Kabel USB', 'HARUS yang data, bukan charge-only. Cara test: colok ke PC, buka Device Manager, harus muncul COM port baru.'),
    ('Semua Sensor', 'Dipower dari pin 3.3V ESP32. JANGAN power sensor dari jalur 5V.'),
    ('Common Ground', 'GND adaptor, GND ESP32, GND semua sensor, dan GND servo HARUS semua terhubung ke satu titik GND yang sama. Tanpa common ground, sinyal akan kacau.'),
    ('DHT11', 'Beli yang MODULE 3-pin (sudah ada PCB + pull-up resistor). Jangan beli yang komponen telanjang 4-pin.'),
]

for label, desc in notes:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

# ========== PIN ASSIGNMENT ==========
doc.add_paragraph('')
heading('PIN ASSIGNMENT - ESP32-S3-N16R8', level=1)

summary_table(
    ['Pin ESP32-S3', 'Terhubung Ke', 'Tipe Sinyal', 'Keterangan'],
    [
        ['GPIO8 (SDA)', 'BH1750 SDA', 'I2C Data', 'Default I2C SDA pada ESP32-S3'],
        ['GPIO9 (SCL)', 'BH1750 SCL', 'I2C Clock', 'Default I2C SCL pada ESP32-S3'],
        ['GPIO4', 'DHT11 DATA', 'Digital', 'Pin data sensor suhu/kelembapan'],
        ['GPIO5', 'Rain Sensor DO', 'Digital Input', 'Output digital sensor hujan'],
        ['GPIO13', 'Servo MG996R Signal (oranye)', 'PWM', 'Sinyal kontrol posisi servo'],
        ['GPIO2', 'LED Warning via R 220 Ohm', 'Digital Output', 'Indikator anti-oven warning'],
        ['VIN', 'Adaptor +5V', 'Power Input', 'Input power utama dari adaptor'],
        ['3.3V', 'VCC semua sensor', 'Power Output', 'Sumber daya untuk BH1750, DHT11, Rain'],
        ['GND', 'Common Ground', 'Ground', 'Semua GND harus terhubung'],
    ]
)

# ========== WIRING DIAGRAM ==========
heading('DIAGRAM WIRING (Teks)', level=1)

wiring = """ADAPTOR DC 5V 3A (colok ke stopkontak 220V)
    |
    v
JACK DC FEMALE SCREW TERMINAL
    |             |
    +5V           GND
    |             |
    |--- VIN ESP32-S3-N16R8
    |             |--- GND ESP32 --- GND BUS (semua GND terhubung di sini)
    |             |                    |--- GND BH1750
    |             |                    |--- GND DHT11
    |             |                    |--- GND Rain Module
    |             |                    |--- GND Servo (coklat)
    |             |                    |--- GND LED
    |             |
    |             |--- 3.3V ESP32 --- 3.3V BUS
    |                                  |--- VCC BH1750
    |                                  |--- VCC DHT11
    |                                  |--- VCC Rain Module
    |
    |--- [CAPACITOR 470uF] --- VCC Servo MG996R (merah)
    |    kaki (+) ke +5V        (power servo langsung dari adaptor,
    |    kaki (-) ke GND         BUKAN dari ESP32)
    |
    v
ESP32-S3-N16R8 GPIO Connections:
    GPIO8  (SDA) -----> BH1750 SDA
    GPIO9  (SCL) -----> BH1750 SCL
    GPIO4  -----> DHT11 DATA
    GPIO5  -----> Rain Sensor DO (JANGAN sambung AO)
    GPIO13 -----> Servo MG996R Signal (oranye)
    GPIO2  -----> Resistor 220 Ohm -----> LED Merah (+) -----> GND"""

p = doc.add_paragraph()
r = p.add_run(wiring)
r.font.name = 'Consolas'
r.font.size = Pt(9)

# ========== FOOTER ==========
doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Smart Greenhouse IoT - Dokumen Komponen v1.0\n13525004@std.stei.itb.ac.id | April 2026')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.italic = True

# ========== SAVE ==========
path = '/home/ninalover/Documents/prd/Daftar_Komponen_Smart_Greenhouse.docx'
doc.save(path)
print(f'Saved: {path}')
